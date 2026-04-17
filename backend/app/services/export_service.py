"""
export_service.py
-----------------
Generate file PDF atau DOCX dari hasil pipeline (narasi + stratkom).
"""

import io
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _now_str() -> str:
    return datetime.now().strftime("%d %B %Y, %H:%M WIB")


# ── PDF ──────────────────────────────────────────────────────────────────────

def generate_pdf(session_data: dict, content_type: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "KPMTitle",
        parent=styles["Title"],
        fontSize=16,
        textColor=colors.HexColor("#1a3c6b"),
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "KPMHeading",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#1a3c6b"),
        spaceBefore=12,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "KPMBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=16,
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        "KPMMeta",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.grey,
        spaceAfter=12,
    )

    story = []

    # Header
    story.append(Paragraph("Dokumen Komunikasi Publik", title_style))
    story.append(Paragraph(
        f"Diterbitkan: {_now_str()} &nbsp;|&nbsp; Session: {session_data.get('session_id', '-')}",
        meta_style,
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1a3c6b")))
    story.append(Spacer(1, 0.3 * cm))

    narasi = session_data.get("narasi") or {}
    stratkom = session_data.get("stratkom") or {}
    regulasi = session_data.get("regulasi") or []

    if content_type in ("narasi", "draft"):
        story.append(Paragraph("Isu", heading_style))
        story.append(Paragraph(narasi.get("isu", "-"), body_style))

        story.append(Paragraph("Narasi Isu", heading_style))
        for line in narasi.get("narasi", "").splitlines():
            if line.strip():
                story.append(Paragraph(line.strip(), body_style))

        if narasi.get("key_points"):
            story.append(Paragraph("Poin Kunci", heading_style))
            for i, pt in enumerate(narasi["key_points"], 1):
                story.append(Paragraph(f"{i}. {pt}", body_style))

    if content_type in ("stratkom", "draft"):
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("Strategi Komunikasi", heading_style))
        for line in stratkom.get("strategi", "").splitlines():
            if line.strip():
                story.append(Paragraph(line.strip(), body_style))

        if stratkom.get("rekomendasi"):
            story.append(Paragraph("Rekomendasi", heading_style))
            for i, r in enumerate(stratkom["rekomendasi"], 1):
                story.append(Paragraph(f"{i}. {r}", body_style))

    if regulasi:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph("Dasar Hukum", heading_style))
        table_data = [["No.", "Nomor Regulasi", "Judul", "Tahun"]]
        for reg in regulasi:
            table_data.append([
                str(regulasi.index(reg) + 1),
                reg.get("nomor", "-"),
                reg.get("judul", "-"),
                str(reg.get("tahun", "-")),
            ])
        tbl = Table(table_data, colWidths=[1 * cm, 4 * cm, 9 * cm, 2 * cm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a3c6b")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(tbl)

    doc.build(story)
    return buf.getvalue()


# ── DOCX ─────────────────────────────────────────────────────────────────────

def generate_docx(session_data: dict, content_type: str) -> bytes:
    doc = Document()

    # Margin
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def add_heading(text: str, level: int = 1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)

    def add_body(text: str):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)

    # Title
    title = doc.add_heading("Dokumen Komunikasi Publik", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)

    meta = doc.add_paragraph(
        f"Diterbitkan: {_now_str()}    |    Session: {session_data.get('session_id', '-')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    narasi = session_data.get("narasi") or {}
    stratkom = session_data.get("stratkom") or {}
    regulasi = session_data.get("regulasi") or []

    if content_type in ("narasi", "draft"):
        add_heading("Isu", level=1)
        add_body(narasi.get("isu", "-"))

        add_heading("Narasi Isu", level=1)
        for line in narasi.get("narasi", "").splitlines():
            if line.strip():
                add_body(line.strip())

        if narasi.get("key_points"):
            add_heading("Poin Kunci", level=2)
            for i, pt in enumerate(narasi["key_points"], 1):
                add_body(f"{i}. {pt}")

    if content_type in ("stratkom", "draft"):
        add_heading("Strategi Komunikasi", level=1)
        for line in stratkom.get("strategi", "").splitlines():
            if line.strip():
                add_body(line.strip())

        if stratkom.get("rekomendasi"):
            add_heading("Rekomendasi", level=2)
            for i, r in enumerate(stratkom["rekomendasi"], 1):
                add_body(f"{i}. {r}")

    if regulasi:
        add_heading("Dasar Hukum", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, label in enumerate(["No.", "Nomor Regulasi", "Judul", "Tahun"]):
            hdr[i].text = label
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for idx, reg in enumerate(regulasi):
            row = tbl.add_row().cells
            row[0].text = str(idx + 1)
            row[1].text = reg.get("nomor", "-")
            row[2].text = reg.get("judul", "-")
            row[3].text = str(reg.get("tahun", "-"))
            for cell in row:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
